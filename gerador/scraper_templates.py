import os
import json
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CONFIGURAÇÕES
# ============================================================
PASTA_REVISTAS = "C:/SubmetNEV/revistas"
PASTA_TEMPLATES = "C:/SubmetNEV/templates"
PASTA_IMAGEM = "C:/SubmetNEV/imagem"
ICONE_PATH = os.path.join(PASTA_IMAGEM, "icone_submetnev.png")
SOBRESCREVER = True

# ============================================================
# FUNÇÕES AUXILIARES DE FORMATAÇÃO (CORES, BORDAS)
# ============================================================
def aplicar_fundo_aviso(paragraph, cor_rgb=(255, 242, 204)):
    """Aplica fundo amarelo claro a um parágrafo inteiro."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), f"{cor_rgb[0]:02X}{cor_rgb[1]:02X}{cor_rgb[2]:02X}")
    paragraph._element.get_or_add_pPr().append(shading_elm)

def aplicar_borda_aviso(paragraph, cor_rgb=(255, 145, 77)):
    """Adiciona borda inferior laranja ao parágrafo."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), f"{cor_rgb[0]:02X}{cor_rgb[1]:02X}{cor_rgb[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)

def adicionar_instrucao(doc, mensagem):
    """
    Adiciona um parágrafo de instrução com fundo amarelo, borda laranja e ícone de dica.
    (Não precisa ser apagado pelo autor – serve como orientação permanente.)
    """
    p = doc.add_paragraph()
    aplicar_fundo_aviso(p)
    aplicar_borda_aviso(p)
    run = p.add_run(f"💡 {mensagem}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    return p

def adicionar_placeholder(doc, texto_placeholder):
    """
    Adiciona um parágrafo NORMAL (sem fundo/borda, sem itálico) com o texto do placeholder.
    O autor deve substituí-lo pelo conteúdo real; a formatação será a do periódico.
    """
    p = doc.add_paragraph()
    run = p.add_run(texto_placeholder)
    # Nenhuma formatação especial: fonte, tamanho, espaçamento serão os do estilo Normal
    return p

# ============================================================
# FUNÇÕES DE FORMATAÇÃO DO DOCUMENTO
# ============================================================
def configurar_pagina(doc, revista_data):
    """Configura margens e orientação com base no JSON ou valores padrão"""
    formatacao = revista_data.get('formatacao', {})
    margens_str = formatacao.get('margens', '2,5 cm')
    
    if margens_str is None or not isinstance(margens_str, str):
        margens_str = '2,5 cm'
    
    numeros = re.findall(r'(\d+[,.]?\d*)', margens_str)
    if numeros:
        margem_cm = float(numeros[0].replace(',', '.'))
    else:
        margem_cm = 2.5
    
    section = doc.sections[0]
    section.top_margin = Cm(margem_cm)
    section.bottom_margin = Cm(margem_cm)
    section.left_margin = Cm(margem_cm)
    section.right_margin = Cm(margem_cm)

def configurar_estilo_normal(doc, revista_data):
    """Aplica estilo Normal com fonte, tamanho, espaçamento e recuo de primeira linha"""
    formatacao = revista_data.get('formatacao', {})
    fonte = formatacao.get('fonte', 'Times New Roman')
    
    tamanho = formatacao.get('tamanho_fonte')
    if tamanho is None or tamanho == "":
        tamanho = 12
    else:
        try:
            tamanho = int(tamanho)
        except (ValueError, TypeError):
            tamanho = 12
    
    espacamento = formatacao.get('espacamento_entrelinhas')
    if espacamento is None or espacamento == "":
        espacamento = 1.5
    elif isinstance(espacamento, str):
        if espacamento.lower() == "simples":
            espacamento = 1
        elif espacamento.lower() == "duplo":
            espacamento = 2
        else:
            try:
                espacamento = float(espacamento)
            except ValueError:
                espacamento = 1.5
    elif isinstance(espacamento, (int, float)):
        pass
    else:
        espacamento = 1.5
    
    recuo = formatacao.get('recuo_primeira_linha', '1,25 cm')
    
    style = doc.styles['Normal']
    style.font.name = fonte
    style.font.size = Pt(tamanho)
    
    if isinstance(espacamento, (int, float)):
        style.paragraph_format.line_spacing = espacamento
    else:
        style.paragraph_format.line_spacing = 1.5
    
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if recuo and isinstance(recuo, str):
        num_recuo = re.search(r'(\d+[,.]?\d*)', recuo)
        if num_recuo:
            recuo_cm = float(num_recuo.group(1).replace(',', '.'))
            style.paragraph_format.first_line_indent = Cm(recuo_cm)

def configurar_estilos_titulos(doc, revista_data):
    """Ajusta títulos (Heading 1, 2, 3) com a mesma fonte da revista e tamanhos relativos"""
    formatacao = revista_data.get('formatacao', {})
    fonte = formatacao.get('fonte', 'Times New Roman')
    
    tamanho_base = formatacao.get('tamanho_fonte')
    if tamanho_base is None or tamanho_base == "":
        tamanho_base = 12
    else:
        try:
            tamanho_base = int(tamanho_base)
        except (ValueError, TypeError):
            tamanho_base = 12
    
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        style.font.name = fonte
        style.font.size = Pt(tamanho_base + 2 - i)
        style.font.bold = True

def adicionar_cabecalho_com_icone(doc, revista_nome):
    """Adiciona cabeçalho com ícone do SubmetNEV e título da revista."""
    section = doc.sections[0]
    header = section.header
    for para in header.paragraphs:
        para.clear()
    
    tabela = header.add_table(rows=1, cols=2, width=Inches(6.5))
    tabela.autofit = False
    tabela.columns[0].width = Inches(1.2)
    tabela.columns[1].width = Inches(5.3)
    
    celula_img = tabela.cell(0, 0)
    celula_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if os.path.exists(ICONE_PATH):
        paragrafo_img = celula_img.paragraphs[0]
        run = paragrafo_img.add_run()
        try:
            run.add_picture(ICONE_PATH, width=Cm(1.5))
        except Exception as e:
            print(f"   ⚠️ Não foi possível inserir ícone: {e}")
    else:
        celula_img.text = "[logo]"
    
    celula_texto = tabela.cell(0, 1)
    celula_texto.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = celula_texto.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"SubmetNEV – {revista_nome}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

def adicionar_rodape_com_marca(doc):
    """Adiciona rodapé com a logomarca e mensagem de copyright."""
    section = doc.sections[0]
    footer = section.footer
    for para in footer.paragraphs:
        para.clear()
    
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(ICONE_PATH):
        run = p.add_run()
        try:
            run.add_picture(ICONE_PATH, width=Cm(1.0))
        except:
            pass
        p.add_run(" ")
    run = p.add_run("Template gerado automaticamente pelo SubmetNEV – NEV/USP")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

def adicionar_numero_pagina_rodape(doc):
    """Adiciona número da página no rodapé (alinhado à direita)"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# ============================================================
# FUNÇÕES DE CONTEÚDO (COM INSTRUÇÕES E PLACEHOLDERS NORMAIS)
# ============================================================
def adicionar_placeholders_titulo_autores(doc, revista_nome):
    """Instrução destacada + placeholder sem formatação especial."""
    doc.add_paragraph().add_run(f"Template para submissão à {revista_nome}").bold = True
    doc.add_paragraph("Este documento contém formatação pré-definida e orientações.")
    doc.add_paragraph()
    
    adicionar_instrucao(doc, "Substitua pelo título do seu manuscrito")
    adicionar_placeholder(doc, "[INSIRA O TÍTULO COMPLETO, CONFORME NORMAS DA REVISTA]")
    
    adicionar_instrucao(doc, "Liste todos os autores com suas respectivas informações (ORCID, afiliação)")
    adicionar_placeholder(doc, "[NOME COMPLETO DOS AUTORES, ORCID, AFILIAÇÃO INSTITUCIONAL]")
    
    adicionar_instrucao(doc, "Indique o endereço eletrônico principal para correspondência")
    adicionar_placeholder(doc, "[E-MAIL DO AUTOR PARA CORRESPONDÊNCIA]")
    
    doc.add_paragraph()

def adicionar_secao_resumo_palavras_chave(doc, limites, revista):
    """Instrução destacada + placeholder sem formatação especial."""
    doc.add_heading("RESUMO", level=1)
    limite_resumo = limites.get('resumo_maximo', '200 palavras')
    if limite_resumo is None:
        limite_resumo = '200 palavras'
    p = doc.add_paragraph()
    p.add_run(f"⚠️ ATENÇÃO: O resumo deve ter no máximo {limite_resumo}. ")
    p.add_run("Estrutura sugerida: contexto, objetivo, método, principais resultados, conclusão.\n\n")
    
    adicionar_instrucao(doc, "Escreva um resumo claro e conciso, dentro do limite de palavras")
    adicionar_placeholder(doc, "[INSIRA O RESUMO AQUI]")
    
    doc.add_heading("Palavras-chave", level=2)
    qtd = limites.get('palavras_chave_quantidade', '3 a 5')
    if qtd is None:
        qtd = '3 a 5'
    doc.add_paragraph(f"Quantidade: {qtd}. Termos descritores do conteúdo.")
    adicionar_instrucao(doc, "Separe as palavras-chave por vírgula ou ponto e vírgula")
    adicionar_placeholder(doc, "[palavra1, palavra2, palavra3, ...]")
    doc.add_paragraph()

def adicionar_estrutura_texto(doc, tipo_info, revista):
    """Instrução destacada + placeholder para cada seção."""
    estrutura = tipo_info.get('estrutura_sugerida', [])
    if not estrutura:
        tipo_nome = tipo_info.get('tipo', '').lower()
        if "resenha" in tipo_nome:
            estrutura = ["Contextualização da obra", "Análise crítica", "Considerações finais"]
        elif "ensaio" in tipo_nome:
            estrutura = ["Tese central e problematização", "Desenvolvimento argumentativo", "Conclusão e implicações"]
        else:
            estrutura = ["INTRODUÇÃO", "REFERENCIAL TEÓRICO", "METODOLOGIA", "RESULTADOS", "DISCUSSÃO", "CONCLUSÃO", "REFERÊNCIAS"]
    for secao in estrutura:
        if secao:
            doc.add_heading(secao, level=1)
            adicionar_instrucao(doc, f"Consulte as diretrizes da revista e substitua este texto pelo conteúdo real da seção '{secao}'")
            adicionar_placeholder(doc, f"[Desenvolva aqui a seção '{secao}']")
            doc.add_paragraph()

def adicionar_checklist_final(doc, checklist):
    """Adiciona uma página com o checklist de submissão (sem alterações)"""
    if not checklist:
        return
    doc.add_page_break()
    doc.add_heading("✅ CHECKLIST DE SUBMISSÃO", level=1)
    doc.add_paragraph("Antes de enviar, verifique cada item abaixo:")
    for item in checklist:
        if item:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

def adicionar_exemplo_citacoes_longas(doc, revista_data):
    """Adiciona seção de exemplo de citação longa (sem alterações)"""
    formatacao = revista_data.get('formatacao', {})
    cit_longas = formatacao.get('citacoes_longas', {})
    if not cit_longas or not isinstance(cit_longas, dict):
        return
    
    doc.add_heading("Exemplo de Citação Longa", level=2)
    
    recuo_str = cit_longas.get('recuo', '4 cm')
    if recuo_str is None or not isinstance(recuo_str, str):
        recuo_str = '4 cm'
    
    fonte_tamanho = cit_longas.get('fonte', 11)
    if fonte_tamanho is None:
        fonte_tamanho = 11
    else:
        try:
            fonte_tamanho = int(fonte_tamanho)
        except (ValueError, TypeError):
            fonte_tamanho = 11
    
    espacamento = cit_longas.get('espacamento', 'simples')
    if espacamento is None:
        espacamento = 'simples'
    
    p = doc.add_paragraph()
    p.add_run(f"⚠️ Recuo: {recuo_str} | Fonte: {fonte_tamanho} pt | Espaçamento: {espacamento}")
    
    exemplo = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris."
    p = doc.add_paragraph(exemplo)
    
    recuo_cm = 4.0
    num_recuo = re.search(r'(\d+[,.]?\d*)', recuo_str)
    if num_recuo:
        recuo_cm = float(num_recuo.group(1).replace(',', '.'))
    p.paragraph_format.left_indent = Cm(recuo_cm)
    
    try:
        p.runs[0].font.size = Pt(fonte_tamanho)
    except (ValueError, TypeError, IndexError):
        p.runs[0].font.size = Pt(11)
    
    if espacamento == 'simples' or espacamento == 1:
        p.paragraph_format.line_spacing = 1
    elif espacamento == 'duplo' or espacamento == 2:
        p.paragraph_format.line_spacing = 2
    else:
        p.paragraph_format.line_spacing = 1

def adicionar_exemplos_citacoes_referencias(doc, revista_data):
    """Adiciona exemplos de citações no texto e referências (sem alterações)"""
    formatacao = revista_data.get('formatacao', {})
    ex_citacoes = formatacao.get('exemplos_citacoes', '')
    ex_refs = formatacao.get('exemplos_referencias', '')
    
    if ex_citacoes and isinstance(ex_citacoes, str) and ex_citacoes.strip():
        doc.add_heading("Exemplo de Citação no Texto (autor-data)", level=2)
        doc.add_paragraph(ex_citacoes)
    
    if ex_refs and isinstance(ex_refs, str) and ex_refs.strip():
        doc.add_heading("Exemplos de Referências Bibliográficas", level=2)
        doc.add_paragraph(ex_refs)

def adicionar_instrucoes_adicionais(doc, instrucoes):
    """Insere orientações extras como declaração de dados, conflito de interesses, etc. (sem alterações)"""
    if not instrucoes:
        return
    doc.add_heading("INSTRUÇÕES ADICIONAIS (conforme diretrizes da revista)", level=1)
    for key, value in instrucoes.items():
        if value and key not in ['preprint']:
            p = doc.add_paragraph()
            p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
            p.add_run(str(value))

# ============================================================
# FUNÇÃO PRINCIPAL: GERAR TEMPLATE PARA UM TIPO DE MANUSCRITO
# ============================================================
def gerar_template(revista_data, tipo_info, caminho_template):
    """
    revista_data: dicionário completo da revista (já carregado)
    tipo_info: dicionário do tipo_texto (contém 'tipo', 'template', 'extensao', etc.)
    caminho_template: onde salvar o .docx
    """
    print(f"   📝 Gerando template: {caminho_template}")
    
    doc = Document()
    
    # Cabeçalho e rodapé
    nome_revista = revista_data.get('nome', 'Revista')
    if nome_revista is None:
        nome_revista = 'Revista'
    adicionar_cabecalho_com_icone(doc, nome_revista)
    adicionar_rodape_com_marca(doc)
    adicionar_numero_pagina_rodape(doc)
    
    # Formatação geral
    configurar_pagina(doc, revista_data)
    configurar_estilo_normal(doc, revista_data)
    configurar_estilos_titulos(doc, revista_data)
    
    # Placeholders com instruções em destaque
    adicionar_placeholders_titulo_autores(doc, nome_revista)
    
    # Resumo e palavras-chave
    limites = revista_data.get('limites', {})
    if limites is None:
        limites = {}
    adicionar_secao_resumo_palavras_chave(doc, limites, revista_data)
    
    # Estrutura do texto
    adicionar_estrutura_texto(doc, tipo_info, revista_data)
    
    # Checklist
    checklist = revista_data.get('checklist', [])
    if checklist is None:
        checklist = []
    adicionar_checklist_final(doc, checklist)
    
    # Exemplos de citações
    adicionar_exemplo_citacoes_longas(doc, revista_data)
    adicionar_exemplos_citacoes_referencias(doc, revista_data)
    
    # Instruções adicionais
    instrucoes = revista_data.get('instrucoes_adicionais', {})
    if instrucoes is None:
        instrucoes = {}
    adicionar_instrucoes_adicionais(doc, instrucoes)
    
    # Aviso de extensão
    extensao = tipo_info.get('extensao')
    if extensao and isinstance(extensao, str) and extensao.strip():
        doc.add_paragraph()
        nota = doc.add_paragraph()
        nota.add_run(f"⚠️ LIMITE DE EXTENSÃO: {extensao}. Respeite esse limite para submissão.").italic = True
    
    os.makedirs(os.path.dirname(caminho_template), exist_ok=True)
    doc.save(caminho_template)

# ============================================================
# LEITURA DE TODOS OS JSONS E GERAÇÃO EM MASSA
# ============================================================
def processar_todas_revistas():
    if not os.path.exists(PASTA_REVISTAS):
        print(f"❌ Pasta {PASTA_REVISTAS} não encontrada.")
        return
    
    jsons = [f for f in os.listdir(PASTA_REVISTAS) if f.endswith('.json') and f != 'index.json']
    
    if not jsons:
        print("⚠️ Nenhum JSON de revista encontrado.")
        return
    
    total_templates = 0
    for json_file in jsons:
        caminho_json = os.path.join(PASTA_REVISTAS, json_file)
        with open(caminho_json, 'r', encoding='utf-8') as f:
            revista = json.load(f)
        
        nome_revista = revista.get('nome', 'Desconhecida')
        if nome_revista is None:
            nome_revista = 'Desconhecida'
        id_rev = revista.get('id', '?')
        print(f"\n📘 Processando revista: {nome_revista} (ID {id_rev})")
        
        tipos = revista.get('tipos_texto', [])
        if not tipos:
            print("   ⚠️ Nenhum tipo_texto encontrado. Pulando.")
            continue
        
        for tipo in tipos:
            caminho_template_rel = tipo.get('template')
            if not caminho_template_rel or caminho_template_rel is None:
                print("   ⚠️ Tipo sem campo 'template'. Pulando.")
                continue
            
            caminho_absoluto = os.path.join(os.path.dirname(PASTA_TEMPLATES), caminho_template_rel)
            
            if os.path.exists(caminho_absoluto) and not SOBRESCREVER:
                print(f"   ⏭️ Template já existe: {caminho_absoluto} (pulando)")
                continue
            
            try:
                gerar_template(revista, tipo, caminho_absoluto)
                total_templates += 1
            except Exception as e:
                print(f"   ❌ Erro ao gerar template {caminho_absoluto}: {e}")
    
    print(f"\n✅ Processamento concluído. {total_templates} template(s) gerado(s) em {PASTA_TEMPLATES}")

if __name__ == "__main__":
    processar_todas_revistas()